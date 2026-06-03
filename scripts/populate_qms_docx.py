"""
populate_qms_docx.py

Populates QMS requirements Word output files from filled markdown files in the output directory.

Requirements template:
  - Each requirement → its own standalone vertical record table (ID / Requirement / Rationale / Category),
    cloned from the template table in the Functional/Feature Requirements section.

Usage:
    python scripts/populate_qms_docx.py <output_dir> <requirements_template.docx>
"""

import sys
import os
import re
import copy
import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def strip_frontmatter(md: str) -> str:
    match = re.match(r"^---\n.*?\n---\n", md, re.DOTALL)
    return md[match.end():].strip() if match else md.strip()


def parse_md_table(text: str) -> list[list[str]]:
    """Parse markdown table rows (skip separator). Returns list of cell lists."""
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if not (line.startswith("|") and line.endswith("|")):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.match(r"^[-: ]+$", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


def extract_section(md: str, heading: str) -> str:
    """Return content under a ## heading (until next same/higher heading)."""
    match = re.search(r"(?m)^##\s+" + re.escape(heading), md, re.IGNORECASE)
    if not match:
        return ""
    start = match.end()
    end_match = re.search(r"(?m)^##\s+", md[start:])
    end = start + end_match.start() if end_match else len(md)
    return md[start:end].strip()


def extract_text_lines(section_content: str) -> list[str]:
    """Return non-table, non-heading lines from a section."""
    lines = []
    for line in section_content.splitlines():
        line = line.strip()
        if not line or line.startswith("|") or line.startswith("#") or re.match(r"^[-*_]{3,}$", line):
            continue
        lines.append(re.sub(r"^\s*[-*]\s*", "• ", line))
    return lines


# ---------------------------------------------------------------------------
# Word XML helpers
# ---------------------------------------------------------------------------

def set_cell_text(cell, text: str):
    """Replace all runs in the cell's first paragraph with a single fresh run."""
    para = cell.paragraphs[0]
    for run_elem in para._p.findall(qn("w:r")):
        para._p.remove(run_elem)
    run = para.add_run(str(text))


def find_heading_element(doc, heading_text: str):
    """Return the XML element of the paragraph matching heading_text (case-insensitive)."""
    for elem in doc.element.body:
        if elem.tag.split("}")[-1] == "p":
            pPr = elem.find(f".//{{{NS}}}pStyle")
            if pPr is not None:
                style = pPr.get(f"{{{NS}}}val", "")
                if "Heading" in style or "heading" in style:
                    text = "".join(t.text or "" for t in elem.iter(f"{{{NS}}}t"))
                    if text.strip().lower() == heading_text.strip().lower():
                        return elem
    return None


def find_table_after_heading(doc, heading_text: str):
    """Return the first Word table that appears after a given heading."""
    found_heading = False
    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1]
        if not found_heading and tag == "p":
            pPr = elem.find(f".//{{{NS}}}pStyle")
            if pPr is not None:
                style = pPr.get(f"{{{NS}}}val", "")
                if "Heading" in style or "heading" in style:
                    text = "".join(t.text or "" for t in elem.iter(f"{{{NS}}}t"))
                    if text.strip().lower() == heading_text.strip().lower():
                        found_heading = True
        elif found_heading and tag == "tbl":
            # Find which index this table is
            all_tbls = doc.element.body.findall(f".//{{{NS}}}tbl")
            idx = list(all_tbls).index(elem)
            return doc.tables[idx]
    return None


def insert_paragraph_after(anchor_elem, text: str, bold: bool = False):
    """Insert a plain paragraph after anchor_elem in the document body."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    anchor_elem.addnext(p)
    return p


def make_table_element(headers: list[str], rows: list[list[str]]) -> etree._Element:
    """Build a minimal <w:tbl> XML element from headers + rows."""
    tbl = OxmlElement("w:tbl")
    # Table properties
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl.append(tblPr)

    def make_row(cells: list[str], header: bool = False) -> etree._Element:
        tr = OxmlElement("w:tr")
        for text in cells:
            tc = OxmlElement("w:tc")
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            if header:
                rPr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rPr.append(b)
                r.append(rPr)
            t = OxmlElement("w:t")
            t.text = str(text)
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        return tr

    tbl.append(make_row(headers, header=True))
    for row in rows:
        tbl.append(make_row(row))
    return tbl


# ---------------------------------------------------------------------------
# Separate-table requirements pattern
# ---------------------------------------------------------------------------

def fill_record_table(table, record: dict, field_keys: list[str]):
    """Fill a single vertical record table (field label col 0, value col 1)."""
    for row_idx, key in enumerate(field_keys):
        if row_idx < len(table.rows) and len(table.rows[row_idx].cells) > 1:
            set_cell_text(table.rows[row_idx].cells[1], record.get(key, ""))


def populate_requirements_as_separate_tables(template_table, records: list[dict], field_keys: list[str]):
    """
    Populate the template table with the first record, then clone it for each
    subsequent record, inserting a paragraph spacer between each table.
    Returns the XML element of the last inserted table.
    """
    if not records:
        return template_table._tbl

    # Save the unfilled template structure before filling req 0
    template_tbl_copy = copy.deepcopy(template_table._tbl)

    # Fill the original template table with the first record
    fill_record_table(template_table, records[0], field_keys)

    anchor = template_table._tbl
    for record in records[1:]:
        # Paragraph spacer
        spacer_p = OxmlElement("w:p")
        anchor.addnext(spacer_p)
        anchor = spacer_p

        # Clone the unfilled template structure so each table starts clean
        new_tbl = copy.deepcopy(template_tbl_copy)
        anchor.addnext(new_tbl)
        anchor = new_tbl

        # Re-wrap as a Document table to use fill_record_table
        from docx.table import Table as DocxTable
        wrapped = DocxTable(new_tbl, template_table._parent)
        fill_record_table(wrapped, record, field_keys)

    return anchor  # XML element of the last inserted table


# ---------------------------------------------------------------------------
# Horizontal table (design template pattern)
# ---------------------------------------------------------------------------

def populate_horizontal_table(table, md_rows: list[list[str]]):
    """
    Fill a horizontal table from md_rows (row 0 = headers, rows 1+ = data).
    Saves the template data row, clears data rows, clones for each incoming row.
    """
    if len(md_rows) < 2:
        return

    num_cols = len(table.columns)
    # Save template data row
    if len(table.rows) > 1:
        template_tr = copy.deepcopy(table.rows[1]._tr)
    else:
        template_tr = copy.deepcopy(table.rows[0]._tr)

    # Clear data rows
    tbl = table._tbl
    for row in table.rows[1:]:
        tbl.remove(row._tr)

    for md_row in md_rows[1:]:
        new_tr = copy.deepcopy(template_tr)
        tbl.append(new_tr)
        new_row = table.rows[-1]
        for col_idx, cell in enumerate(new_row.cells):
            text = md_row[col_idx] if col_idx < len(md_row) else ""
            set_cell_text(cell, text)


# ---------------------------------------------------------------------------
# Requirements document population
# ---------------------------------------------------------------------------

def populate_requirements_doc(md_path: str, template_path: str, docx_out: str):
    with open(md_path, encoding="utf-8") as f:
        md = strip_frontmatter(f.read())

    shutil.copy2(template_path, docx_out)
    doc = Document(docx_out)

    # --- Product Requirements → vertical records in Functional/Feature Requirements (Table 2) ---
    req_section = extract_section(md, "Product Requirements")
    req_rows = parse_md_table(req_section)
    # Expected columns: ID | Requirement | Source artefact | Priority | Rationale
    if len(req_rows) > 1:
        records = []
        for row in req_rows[1:]:  # skip header
            records.append({
                "id":          row[0] if len(row) > 0 else "",
                "requirement": row[1] if len(row) > 1 else "",
                "rationale":   row[4] if len(row) > 4 else "",
                "category":    row[3] if len(row) > 3 else "",
            })
        table = find_table_after_heading(doc, "Functional/Feature Requirements")
        if table:
            last_req_elem = populate_requirements_as_separate_tables(table, records, ["id", "requirement", "rationale", "category"])
        else:
            last_req_elem = None

    # --- Traceability → insert after the last requirement table ---
    trace_section = extract_section(md, "Traceability")
    trace_rows = parse_md_table(trace_section)
    if len(trace_rows) > 1 and last_req_elem is not None:
        tbl_elem = make_table_element(trace_rows[0], trace_rows[1:])
        last_req_elem.addnext(tbl_elem)
        tbl_elem.addprevious(make_heading_element("Traceability", level=2))

    # --- Gaps and observations → paragraph text after Data Requirements heading ---
    gaps_section = extract_section(md, "Gaps and observations")
    gap_lines = extract_text_lines(gaps_section)
    if gap_lines:
        heading_elem = find_heading_element(doc, "Data Requirements")
        if heading_elem is not None:
            anchor = heading_elem
            for line in reversed(gap_lines):
                p = insert_paragraph_after(anchor, line)
            insert_paragraph_after(anchor, "Gaps and observations from source analysis:", bold=True)

    doc.save(docx_out)
    print(f"Saved: {docx_out}")


def make_heading_element(text: str, level: int = 2) -> etree._Element:
    """Create a heading paragraph XML element."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), f"Heading{level}")
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


# ---------------------------------------------------------------------------
# Verdict document population
# ---------------------------------------------------------------------------

def _extract_verdict_summary(md: str) -> dict:
    """Extract winner, score, and scorecard rows from the verdict markdown."""
    summary = {"winner": "", "score": "", "scorecard_rows": []}

    winner_m = re.search(r"\*\*Winner:\*\*\s*(.+)", md)
    if winner_m:
        summary["winner"] = winner_m.group(1).strip()

    score_m = re.search(r"\*\*Score:\*\*\s*(.+)", md)
    if score_m:
        summary["score"] = score_m.group(1).strip()

    scorecard_m = re.search(r"## Per-criterion scorecard(.*?)(?=\n##|\Z)", md, re.DOTALL)
    if scorecard_m:
        summary["scorecard_rows"] = parse_md_table(scorecard_m.group(1))

    return summary


def _add_verdict_banner(doc, summary: dict):
    """Write a prominent verdict summary at the top of the document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_heading("Final Verdict", level=1)

    # Winner line — large bold text
    winner_para = doc.add_paragraph()
    winner_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = winner_para.add_run(f"Winner: {summary['winner']}")
    run.bold = True
    run.font.size = Pt(16)

    # Score line
    score_para = doc.add_paragraph()
    run = score_para.add_run(f"Score: {summary['score']}")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("")

    # Scorecard table
    rows = summary["scorecard_rows"]
    if rows:
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < len(tbl.columns):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(cell_text)
                    if r_idx == 0 or (r_idx == len(rows) - 1):
                        run.bold = True

    doc.add_paragraph("")


def populate_verdict_doc(md_path: str, docx_out: str):
    """Convert the QA verdict markdown to a Word document with headings, tables, and bold text."""
    with open(md_path, encoding="utf-8") as f:
        md = strip_frontmatter(f.read())

    doc = Document()

    # Verdict summary banner at the top
    summary = _extract_verdict_summary(md)
    _add_verdict_banner(doc, summary)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1
            continue

        # Horizontal rule — skip
        if re.match(r"^[-*_]{3,}\s*$", line):
            i += 1
            continue

        # Markdown table — collect all contiguous table lines
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_md_table("\n".join(table_lines))
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < len(tbl.columns):
                            cell = tbl.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            run = cell.paragraphs[0].add_run(cell_text)
                            if r_idx == 0:
                                run.bold = True
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        # Strip HTML comments and skip if empty after stripping
        stripped = re.sub(r"<!--.*?-->", "", line).strip()
        if not stripped:
            i += 1
            continue

        # Regular paragraph — handle **bold** spans
        para = doc.add_paragraph()
        parts = re.split(r"\*\*(.*?)\*\*", stripped)
        for j, part in enumerate(parts):
            run = para.add_run(part)
            if j % 2 == 1:
                run.bold = True
        i += 1

    doc.save(docx_out)
    print(f"Saved: {docx_out}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def populate(output_dir: str, requirements_template: str):
    files = [f for f in os.listdir(output_dir) if f.endswith(".md") and "qms-flow" in f]
    for filename in sorted(files):
        filepath = os.path.join(output_dir, filename)
        docx_out = os.path.join(output_dir, filename.replace(".md", ".docx"))
        if "-requirements" in filename:
            populate_requirements_doc(filepath, requirements_template, docx_out)
        elif "-qa-verdict" in filename:
            populate_verdict_doc(filepath, docx_out)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scripts/populate_qms_docx.py <output_dir> <requirements_template.docx>")
        sys.exit(1)
    populate(sys.argv[1], sys.argv[2])
