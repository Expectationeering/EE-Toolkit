"""
populate_docx.py

Populates a copied Expectationeering Workbook .docx from a filled-in markdown workbook.

The markdown workbook uses horizontal pipe tables for every artifact. In the Word
template, some artifacts are laid out horizontally (header row + data rows) and others
"vertically" (one table per record, with [field label | value] rows). This script maps
each markdown section to the corresponding Word table by a fixed table index, captured
once up front so that inserting record tables does not invalidate later references.

Vertical-record tables are filled POSITIONALLY: the order of the markdown columns matches
the order of the rows in the Word template, so no fragile field-label matching is needed.

Usage:
    python scripts/populate_docx.py <workbook.md> <template.docx> <output.docx>
"""

import re
import os
import glob
import sys
import copy
import shutil
import argparse
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

def extract_section(md: str, heading: str) -> str:
    """Return the markdown content under a heading (until the next same/higher heading)."""
    level = len(re.match(r'^(#+)', heading).group(1)) if re.match(r'^(#+)', heading) else 2
    pattern = re.escape(heading.lstrip('#').strip())
    match = re.search(r'(?m)^#{1,' + str(level) + r'}\s+' + pattern, md, re.IGNORECASE)
    if not match:
        return ''
    start = match.end()
    end_match = re.search(r'(?m)^#{1,' + str(level) + r'}\s+', md[start:])
    end = start + end_match.start() if end_match else len(md)
    return md[start:end].strip()


def parse_md_table(text: str) -> list:
    """Parse a markdown pipe table. Returns list of rows (header first), each a list of cells."""
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line.strip('|').split('|')]
            if all(re.match(r'^[-:]+$', c) for c in cells if c):
                continue
            rows.append(cells)
    return rows


def is_placeholder(value: str) -> bool:
    """A cell that is empty or still holds a template placeholder/comment."""
    v = (value or '').strip()
    if not v:
        return True
    if v.startswith('<!--') or (v.startswith('<') and v.endswith('>')):
        return True
    # bare ID prefix like 'UE_' with nothing after it
    if re.fullmatch(r'[A-Z_]+_', v):
        return True
    return False


def data_rows(md_rows: list) -> list:
    """Return the real data rows of a parsed table (skip header, drop placeholder-only rows)."""
    if not md_rows or len(md_rows) < 2:
        return []
    out = []
    for r in md_rows[1:]:
        # keep the row if its ID cell (col 0) holds a real value
        if r and not is_placeholder(r[0]):
            out.append(r)
    return out


def parse_feature_records(md: str) -> list:
    """Parse the BDD feature files (```gherkin blocks) in the Verification section into SV records.

    Each feature file becomes one record [SV_id, feature_text, result, traces], where traces is the
    requirement id from the `@ID:...` tag (with any trailing `.n` sub-index stripped).
    """
    section = extract_section(md, '### Verification (SV_*)')
    records = []
    for i, m in enumerate(re.finditer(r'```gherkin\s*\n(.*?)\n```', section, re.DOTALL | re.IGNORECASE), 1):
        feature = m.group(1).strip()
        if not feature:
            continue
        tag = re.search(r'@ID:\s*(\S+)', feature)
        traces = re.sub(r'\.\d+$', '', tag.group(1)) if tag else ''
        records.append([f'SV_{i:02d}', feature, '', traces])
    return records


def get_narrative(md: str, heading: str) -> str:
    """Return the narrative prose under a heading (skipping headings, tables, comments, placeholders)."""
    section = extract_section(md, heading)
    lines = []
    for line in section.splitlines():
        s = line.strip()
        if not s or s.startswith('#') or s.startswith('|') or s.startswith('<!--'):
            continue
        if s.startswith('<') and s.endswith('>'):  # untouched placeholder
            continue
        if re.match(r'^\*\*[^*]+\*\*\s*:', s):  # bold field line, e.g. **Stakeholder**: ...
            continue
        lines.append(s.replace('**', ''))  # drop markdown bold markers
    return '\n'.join(lines).strip()


def replace_placeholder_paragraph(doc, needle: str, text: str) -> bool:
    """Replace the first body paragraph containing `needle` with `text` (preserving formatting)."""
    if not text:
        return False
    for p in doc.paragraphs:
        if needle in p.text:
            parts = text.split('\n')
            for r in list(p.runs):
                r.text = ''
            if p.runs:
                p.runs[0].text = parts[0]
            else:
                p.add_run(parts[0])
            for extra in parts[1:]:
                run = p.add_run()
                run.add_break()
                run.text = extra
            return True
    return False


def derive_input_title(md_path: str) -> str:
    """Return the source input document's title — the first '# ' heading of the matching inputs/*.md.

    The workbook is named '<stem>_Workbook.md', so the input is 'inputs/<stem>.md'. Falls back to
    any markdown file in inputs/ if the exact match is not found.
    """
    stem = os.path.splitext(os.path.basename(md_path))[0]
    if stem.endswith('_Workbook'):
        stem = stem[:-len('_Workbook')]
    candidates = [os.path.join('inputs', stem + '.md')]
    candidates += sorted(glob.glob(os.path.join('inputs', '*.md')))
    for path in candidates:
        if os.path.exists(path):
            with open(path, encoding='utf-8') as f:
                for line in f:
                    m = re.match(r'^#\s+(.+)$', line.strip())
                    if m:
                        return m.group(1).strip()
    return ''


def _set_para_runs(p, text: str):
    """Replace all runs of a paragraph with a single run of `text`, keeping the first run's format."""
    for r in list(p.runs):
        r.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def _replace_in_paragraph(p, old: str, new: str):
    """Replace `old` with `new` in a paragraph (run-level first, whole-paragraph fallback)."""
    if old not in p.text:
        return
    if any(old in r.text for r in p.runs):
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
    else:
        _set_para_runs(p, p.text.replace(old, new))


def _header_paragraphs(hdr):
    """Yield every paragraph in a header/footer, including those inside its tables."""
    for p in hdr.paragraphs:
        yield p
    for t in hdr.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def fill_product_name(doc, name: str):
    """Fill the product name on the front page (the 'Product Name' title) and in every header."""
    if not name:
        return
    # Front-page title heading
    for p in doc.paragraphs:
        if p.style and p.style.name in ('Heading 1', 'Heading1', 'Title') and p.text.strip() == 'Product Name':
            _set_para_runs(p, name)
            break
    # Headers of every section (first-page, default, even), including header tables
    for section in doc.sections:
        for hdr in (section.first_page_header, section.header, section.even_page_header):
            for p in _header_paragraphs(hdr):
                _replace_in_paragraph(p, '<Product Name>', name)
    # Document Title property
    try:
        doc.core_properties.title = name
    except Exception:
        pass


def get_field(md: str, heading: str, field: str) -> str:
    """Return the value of a `**field**: value` line within a section, or '' if absent/placeholder."""
    section = extract_section(md, heading)
    m = re.search(r'(?im)^\*\*' + re.escape(field) + r'\*\*\s*:\s*(.+)$', section)
    if m:
        val = m.group(1).strip()
        if val and not val.startswith('<!--') and not (val.startswith('<') and val.endswith('>')):
            return val
    return ''


def first_subheading(md: str, heading: str) -> str:
    """Return the first non-placeholder sub-heading text inside a section."""
    for m in re.finditer(r'(?m)^#{4,6}\s+(.+)$', extract_section(md, heading)):
        t = m.group(1).strip()
        if t and not t.startswith('<!--') and not (t.startswith('<') and t.endswith('>')):
            return t
    return ''


def populate_user_groups(table, md_rows: list):
    """Best-effort fill of the User Groups grid: fill cells positionally, clear leftover placeholders."""
    rows = data_rows(md_rows)
    body_rows = table.rows[1:]
    for i, tr in enumerate(body_rows):
        cells = tr.cells
        if i < len(rows):
            for c, val in enumerate(rows[i]):
                if c < len(cells):
                    set_value_cell(cells[c]._tc, val)
        else:
            for cell in cells:
                if '<' in cell.text and '>' in cell.text:
                    set_value_cell(cell._tc, '')


# ---------------------------------------------------------------------------
# Word cell / row helpers
# ---------------------------------------------------------------------------

# BDD feature files in the Verification section render in a fixed monospace font so the
# Gherkin indentation and data-table columns stay aligned (see flows/.../Example.feature).
FEATURE_FONT = 'Consolas'
FEATURE_FONT_HALF_PT = '18'  # Word sizes are in half-points; 18 = 9 pt


def _apply_run_font(r_el, name: str, half_pt: str):
    """Set the font name and size (in half-points) on a Word run element `r_el`."""
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)
    rpr.append(rfonts)
    for tag in ('w:sz', 'w:szCs'):
        sz = OxmlElement(tag)
        sz.set(qn('w:val'), half_pt)
        rpr.append(sz)
    r_el.insert(0, rpr)  # rPr must be the first child of the run


def set_value_cell(cell_el, text: str):
    """Set a Word table cell element to a single run of `text`, replacing ALL its paragraphs.

    Feature-file content (a Gherkin block, recognised by its leading `@ID:` tag) is rendered in
    the fixed monospace FEATURE_FONT/size so its indentation and data tables stay aligned.
    """
    text = (text or '').replace('**', '')  # drop markdown bold markers; cells are plain text
    is_feature = text.lstrip().startswith('@ID:')
    paras = cell_el.findall(qn('w:p'))
    # Drop any extra paragraphs so multi-line placeholder text is fully replaced.
    for extra in paras[1:]:
        cell_el.remove(extra)
    p = paras[0] if paras else None
    if p is None:
        p = OxmlElement('w:p')
        cell_el.append(p)
    for r in p.findall(qn('w:r')):
        p.remove(r)
    if text:
        r_new = OxmlElement('w:r')
        for idx, line in enumerate(text.split('\n')):
            if idx > 0:
                r_new.append(OxmlElement('w:br'))  # line break inside the cell
            t_new = OxmlElement('w:t')
            t_new.text = line
            t_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_new.append(t_new)
        if is_feature:
            _apply_run_font(r_new, FEATURE_FONT, FEATURE_FONT_HALF_PT)
        p.append(r_new)


def add_row_copy(table, template_tr, values: list):
    """Deep-copy a horizontal template row and fill its cells positionally."""
    new_tr = copy.deepcopy(template_tr)
    cells = new_tr.findall('.//' + qn('w:tc'))
    for i, val in enumerate(values):
        if i < len(cells):
            set_value_cell(cells[i], val)
    table._tbl.append(new_tr)


def populate_horizontal_table(table, md_rows: list):
    """Replace the data rows of a horizontal Word table with the markdown data rows."""
    rows = data_rows(md_rows)
    if not rows or len(table.rows) < 2:
        return
    template_tr = copy.deepcopy(table.rows[1]._tr)
    for row in table.rows[1:]:
        table._tbl.remove(row._tr)
    for row_data in rows:
        add_row_copy(table, template_tr, row_data)


# ---------------------------------------------------------------------------
# Vertical-record tables (one table per record, filled positionally)
# ---------------------------------------------------------------------------

def _fill_positional_trs(template_trs, values: list):
    """Fill a copy of the template rows: row i's value cell (col 1) gets values[i]."""
    out = []
    for i, tr in enumerate(template_trs):
        new_tr = copy.deepcopy(tr)
        cells = new_tr.findall('.//' + qn('w:tc'))
        if len(cells) >= 2 and i < len(values):
            set_value_cell(cells[1], values[i])
        out.append(new_tr)
    return out


def populate_vertical_positional(table, records: list):
    """
    Fill a vertical-layout section. `records` is a list of value-lists, each value-list being
    one markdown data row (in column order, ID first). Every record becomes its OWN table,
    separated by a blank paragraph, so Word never merges consecutive records into one table.
    """
    if not records or len(table.rows) < 1:
        return
    template_trs = [copy.deepcopy(row._tr) for row in table.rows]
    tbl = table._tbl

    empty_tbl = copy.deepcopy(tbl)
    for tr in empty_tbl.findall(qn('w:tr')):
        empty_tbl.remove(tr)

    for row in list(table.rows):
        tbl.remove(row._tr)
    for new_tr in _fill_positional_trs(template_trs, records[0]):
        tbl.append(new_tr)

    last_tbl = tbl
    for rec in records[1:]:
        new_tbl = copy.deepcopy(empty_tbl)
        for new_tr in _fill_positional_trs(template_trs, rec):
            new_tbl.append(new_tr)
        spacer_p = OxmlElement('w:p')
        last_tbl.addnext(spacer_p)
        spacer_p.addnext(new_tbl)
        last_tbl = new_tbl


def remove_table_and_heading(table):
    """Remove a Word table element and the heading paragraph directly above it."""
    tbl = table._tbl
    prev = tbl.getprevious()
    if prev is not None and prev.tag == qn('w:p'):
        prev.getparent().remove(prev)
    tbl.getparent().remove(tbl)


def clear_heading_text(table):
    """Blank the placeholder text of the heading paragraph directly above a table."""
    prev = table._tbl.getprevious()
    if prev is not None and prev.tag == qn('w:p'):
        for r in prev.findall(qn('w:r')):
            prev.remove(r)


# ---------------------------------------------------------------------------
# Use Cases (best-effort label fill of the complex UC table)
# ---------------------------------------------------------------------------

UC_LABEL_TO_COL = {
    'Title': 'Title', 'Actor': 'Actor', 'Goal': 'Goal', 'Satisfies': 'Satisfies',
    'Classification': 'Classification', 'Precondition': 'Precondition',
    'Post Condition': 'Post Condition', 'Traces': 'Traces',
}


def fill_use_case_table(tbl_el, record: dict):
    """Fill the scalar fields of one UC table element by matching the col-0 label."""
    for tr in tbl_el.findall('.//' + qn('w:tr')):
        cells = tr.findall('.//' + qn('w:tc'))
        if len(cells) < 2:
            continue
        label = ''.join(t.text or '' for t in cells[0].findall('.//' + qn('w:t'))).strip()
        if label == 'ID':
            set_value_cell(cells[1], record.get('ID', ''))
        elif label in UC_LABEL_TO_COL:
            set_value_cell(cells[1], record.get(UC_LABEL_TO_COL[label], ''))


def populate_use_cases(table, records: list):
    """Fill the first UC table and clone a labelled copy for each further use case."""
    if not records:
        return
    template_tbl = copy.deepcopy(table._tbl)
    fill_use_case_table(table._tbl, records[0])
    last_tbl = table._tbl
    for rec in records[1:]:
        new_tbl = copy.deepcopy(template_tbl)
        fill_use_case_table(new_tbl, rec)
        spacer_p = OxmlElement('w:p')
        last_tbl.addnext(spacer_p)
        spacer_p.addnext(new_tbl)
        last_tbl = new_tbl


def records_from_columns(md_rows: list, columns: list) -> list:
    """Build a list of dicts from data rows, keyed by the given column names."""
    out = []
    for r in data_rows(md_rows):
        padded = r + [''] * len(columns)
        out.append({col: padded[i] for i, col in enumerate(columns)})
    return out


# ---------------------------------------------------------------------------
# Main population
# ---------------------------------------------------------------------------

# Table indices in the template (captured by inspection of Expectationeering-Workbook.docx).
T = {
    'DC': 0, 'UE': 1, 'ME': 2, 'BE': 3, 'RE': 4, 'KA': 5, 'BR': 6,
    'IU': 7, 'MD': 8, 'SOI': 9, 'CTX': 10, 'IF': 11, 'USERGROUPS': 12,
    'UR1': 13, 'UR2': 14, 'USER_DFMEA': 15, 'UT': 16, 'UFMEA': 17, 'USR': 18,
    'ACTORS': 19, 'UC': 20, 'DD': 21, 'RQ_IF': 22, 'RQ_FN1': 23, 'RQ_FN2': 24,
    'RQ_PR': 25, 'RQ_NF': 26, 'RQ_CS': 27, 'SV': 28,
}


def populate(md_path: str, template_path: str, output_path: str):
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    with open(md_path, encoding='utf-8') as f:
        md = f.read()

    tables = doc.tables  # cached once; entries are stable references

    # --- Product title: taken from the INPUT document's title; shown on front page + headers ---
    title = derive_input_title(md_path)
    if not title:
        prod = re.search(r'(?im)^\*\*Product\*\*\s*:\s*(.+)$', md)
        title = prod.group(1).strip() if prod else ''
    fill_product_name(doc, title)

    # --- Narrative prose sections (placeholder paragraphs, not tables) ---
    narratives = [
        ('#### Domain Description', 'Describe the applicable domain'),
        ('#### Actual State', 'Describe the pros and cons of the actual'),
        ('#### Desired State', 'Describe the pros and cons of the desired'),
        ('## Intended users of the document', 'Users in scope'),
        ('## Scope of the document', 'Product in scope'),
        ('#### Product Information', 'Description of the Product'),
        ('#### User Expectations (UE_*)', 'Description of User Stakeholder'),
        ('#### Market Expectations (ME_*)', 'Description of Market Stakeholder'),
        ('#### Business Expectations (BE_*)', 'Description of Business Stakeholder'),
        ('#### Regulatory Expectations (RE_*)', 'Description of Regulatory Stakeholder'),
    ]
    for heading, needle in narratives:
        replace_placeholder_paragraph(doc, needle, get_narrative(md, heading))

    # --- Name headings: stakeholder names, solution name, first scenario / use-case titles ---
    for heading, needle in [
        ('#### User Expectations (UE_*)', '<User Stakeholder>'),
        ('#### Market Expectations (ME_*)', '<Market Stakeholder>'),
        ('#### Business Expectations (BE_*)', '<Business Stakeholder>'),
        ('#### Regulatory Expectations (RE_*)', '<Regulatory Stakeholder>'),
    ]:
        replace_placeholder_paragraph(doc, needle, get_field(md, heading, 'Stakeholder'))

    sol = re.search(r'(?im)^##\s*SOLUTION:\s*(.+)$', md)
    if sol:
        replace_placeholder_paragraph(doc, 'SOLUTION: <NAME>', 'SOLUTION: ' + sol.group(1).strip())

    replace_placeholder_paragraph(doc, '<Scenario 1>', first_subheading(md, '### Use Scenarios'))

    uc_data = data_rows(parse_md_table(extract_section(md, '### Use Cases (UC_*)')))
    if uc_data and len(uc_data[0]) > 1:
        replace_placeholder_paragraph(doc, '<Use-Case 1>', uc_data[0][1])

    # --- Context diagram: graphical rendering removed; leave a "To be added" placeholder ---
    replace_placeholder_paragraph(doc, 'Add graphical image', 'To be added')

    def horiz(key, heading):
        populate_horizontal_table(tables[T[key]], parse_md_table(extract_section(md, heading)))

    def vert(key, heading):
        populate_vertical_positional(tables[T[key]], data_rows(parse_md_table(extract_section(md, heading))))

    # --- Application: horizontal tables ---
    horiz('DC', '#### Identified Gaps (DC_*)')
    horiz('UE', '#### User Expectations (UE_*)')
    horiz('ME', '#### Market Expectations (ME_*)')
    horiz('BE', '#### Business Expectations (BE_*)')
    horiz('RE', '#### Regulatory Expectations (RE_*)')
    horiz('KA', "### Ideal Product Model (KA_*)")

    # --- Application: vertical records ---
    vert('BR', "### Business 'Requirements' (BR_*)")

    # --- Context ---
    vert('IU', '### Intended Use (IU_01)')
    vert('MD', '### Medical Device Classification (MD_01)')
    horiz('SOI', '#### System of Interest')
    horiz('CTX', '#### Context Elements')
    horiz('IF', '#### External Interfaces (IF_*)')

    # --- Users ---
    populate_user_groups(tables[T['USERGROUPS']], parse_md_table(extract_section(md, '### User Groups')))
    # User Requirements: two placeholder tables (T13/T14). Fill T13, drop T14.
    ur_records = data_rows(parse_md_table(extract_section(md, '### User Requirements / Needs (UR_*)')))
    if ur_records:
        clear_heading_text(tables[T['UR1']])
        populate_vertical_positional(tables[T['UR1']], ur_records)
    remove_table_and_heading(tables[T['UR2']])
    vert('USER_DFMEA', '### User DFMEA (USER_DFMEA_*)')
    horiz('UT', '### Use Scenarios')
    vert('UFMEA', '### Usability FMEA (UFMEA_*)')
    vert('USR', '### Usability Requirements (USR_*)')

    # --- Concept ---
    horiz('ACTORS', '### Actors')
    uc_records = records_from_columns(
        parse_md_table(extract_section(md, '### Use Cases (UC_*)')),
        ['ID', 'Title', 'Actor', 'Goal', 'Satisfies', 'Classification', 'Precondition',
         'Main Success Scenario', 'Alternative Scenarios', 'Exception Scenarios',
         'Post Condition', 'Traces'],
    )
    populate_use_cases(tables[T['UC']], uc_records)
    vert('DD', '### Design Decisions (DD_*)')

    # --- Development: requirements ---
    vert('RQ_IF', '#### Interface Requirements (RQ_IF_*)')
    # Functional Requirements: two placeholder tables (T23/T24). Fill T23, drop T24.
    fn_records = data_rows(parse_md_table(extract_section(md, '#### Functional Requirements (RQ_FN_*)')))
    if fn_records:
        clear_heading_text(tables[T['RQ_FN1']])
        populate_vertical_positional(tables[T['RQ_FN1']], fn_records)
    remove_table_and_heading(tables[T['RQ_FN2']])
    vert('RQ_PR', '#### Performance Requirements (RQ_PR_*)')
    vert('RQ_NF', '#### Non-Functional Requirements (RQ_NF_*)')
    vert('RQ_CS', '#### Constraint Requirements (RQ_CS_*)')

    # --- Verification: one BDD feature file (Gherkin block) per row ---
    populate_vertical_positional(tables[T['SV']], parse_feature_records(md))

    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Populate Expectationeering Workbook .docx from markdown')
    parser.add_argument('input_md', help='Path to filled-in markdown workbook')
    parser.add_argument('template_docx', help='Path to Expectationeering Workbook .docx template')
    parser.add_argument('output_docx', help='Path for output .docx')
    args = parser.parse_args()
    populate(args.input_md, args.template_docx, args.output_docx)
