"""
md_to_docx.py

Combines all .md files in an output directory into a single Word document.
Markdown headings become Word headings; other lines become paragraphs.

Usage:
    python scripts/md_to_docx.py <output_dir> <output.docx>
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt


def md_to_docx(output_dir: str, output_path: str) -> None:
    md_files = sorted(
        f for f in os.listdir(output_dir) if f.endswith(".md")
    )
    if not md_files:
        print(f"No .md files found in {output_dir}")
        sys.exit(1)

    doc = Document()

    for i, filename in enumerate(md_files):
        if i > 0:
            doc.add_page_break()

        filepath = os.path.join(output_dir, filename)
        with open(filepath, encoding="utf-8") as f:
            lines = f.read().splitlines()

        for line in lines:
            # Headings
            heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=level)
                continue

            # Horizontal rule — skip
            if re.match(r"^[-*_]{3,}\s*$", line):
                continue

            # Blank line — add empty paragraph as spacing
            if not line.strip():
                doc.add_paragraph("")
                continue

            # Regular paragraph (strip leading markdown bold/italic markers for display)
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python scripts/md_to_docx.py <output_dir> <output.docx>")
        sys.exit(1)

    md_to_docx(sys.argv[1], sys.argv[2])
